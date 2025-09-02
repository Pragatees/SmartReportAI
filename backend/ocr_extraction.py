import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import io
import os
import shutil
from fastapi import APIRouter, UploadFile, File
from fastapi.responses import JSONResponse
from docx import Document

# ✅ Detect Tesseract executable path dynamically (Linux-friendly)
tesseract_path = shutil.which("tesseract")
if tesseract_path:
    pytesseract.pytesseract.tesseract_cmd = tesseract_path
else:
    raise RuntimeError("Tesseract is not installed or not in PATH")

router = APIRouter(prefix="/ocr", tags=["OCR"])


@router.post("/extract-text")
async def extract_text(file: UploadFile = File(...)):
    try:
        contents = await file.read()
        file_extension = os.path.splitext(file.filename)[1].lower()
        text = ""

        # 📄 PDF Handling
        if file_extension == ".pdf":
            doc = fitz.open(stream=contents, filetype="pdf")
            for page_num in range(len(doc)):
                page = doc[page_num]
                page_text = page.get_text("text")  # Extract normal text

                if page_text.strip():
                    # ✅ Normal text-based PDF
                    text += page_text + "\n"
                else:
                    # ✅ Scanned PDF (render as image for OCR)
                    pix = page.get_pixmap(dpi=300)  # High DPI for better OCR
                    img = Image.open(io.BytesIO(pix.tobytes("png")))
                    img_text = pytesseract.image_to_string(img)
                    text += img_text + "\n"

                text += f"\n--- End of Page {page_num + 1} ---\n"
            doc.close()

        # 🖼 Image Handling
        elif file_extension in [".png", ".jpg", ".jpeg", ".tif", ".tiff"]:
            img = Image.open(io.BytesIO(contents))
            text = pytesseract.image_to_string(img)

        # 📑 Word Document Handling
        elif file_extension in [".doc", ".docx"]:
            try:
                doc = Document(io.BytesIO(contents))
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text += paragraph.text + "\n"
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text += cell.text + "\n"
            except Exception as doc_e:
                return JSONResponse(
                    content={"error": f"Failed to process Word document: {str(doc_e)}"},
                    status_code=400
                )

        else:
            return JSONResponse(
                content={"error": f"Unsupported file format: {file_extension}"},
                status_code=400
            )

        return {"filename": file.filename, "text": text.strip()}

    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=500)
